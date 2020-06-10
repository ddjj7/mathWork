#coding:utf-8
'''
Created on 2020年2月14日

@author: HEWEI
'''

from docx import Document
from docx.shared import Inches,Pt
from docx.oxml.ns import qn
import random,os




fileseq=0

def chg_font(obj,fontname='微软雅黑',size=None):
    obj.font.name=fontname
    obj._element.rPr.rFonts.set(qn('w:eastAsia'),fontname)
    if size and isinstance(size, Pt):
        obj.font.size=size
        
def checkAndAddFormula(inta,symbol,intb,questline,questlist):
    if symbol=='+':
        questlist.append(questline)
    if symbol=='-':
        if inta>intb:
            questlist.append(questline)
    if symbol=='×':
        if inta!=0 and intb!=0:
            questlist.append(questline)
    if symbol=='÷':
        if inta!=0 and intb!=0 and inta%intb==0:
            questlist.append(questline)
    
        
def genContent(checkv,start,end,num):#count=30,plus=True,minus=True,multiply=False,divide=False):
    allsymbol=['+', '-', '×','÷']
    selectedsymbol=[]
    questlist=[]
    if checkv[0].get()==1:
        selectedsymbol.append(allsymbol[0])
    if checkv[1].get()==1:
        selectedsymbol.append(allsymbol[1])
    if checkv[2].get()==1:
        selectedsymbol.append(allsymbol[2])
    if checkv[3].get()==1:
        selectedsymbol.append(allsymbol[3])
    while True:
        inta=random.randint(start,end)
        stra=str(inta)
        symbol=random.choice(selectedsymbol)
        intb=random.randint(start,end)
        strb=str(intb)
        questline=stra+symbol+strb+'='
        checkAndAddFormula(inta,symbol,intb,questline,questlist)
        if questlist.__len__()==num:
            break
    for s in questlist:
        print(s)
    return questlist

def getFileName(fileseq):
    rootpath='D:\\Program Files\\liClipseWorkspace\\'
    filename=rootpath+'a'+str(fileseq)+'.docx'
    if os.path.exists(filename):
        filename=getFileName(fileseq+1)
    return filename
        
def genDoc(checkv,start,end,num):
    doc=Document()
    distance=Inches(1)
    sec=doc.sections[0]
    sec.left_margin=distance
    sec.right_margin=distance
    sec.top_margin=distance
    sec.bottom_margin=distance
    sec.page_width=Inches(8.27)
    sec.page_height=Inches(11.75)
    
    chg_font(doc.styles['Normal'], fontname='宋体',size=Pt(14))
    
    questlist=genContent(checkv,start,end,num)
    
#     paragraph=doc.add_paragraph('姓名__________    开始时间__________    结束时间__________')
#     ph_format=paragraph.paragraph_format
#     ph_format.space_before=Pt(10)
#     ph_format.space_after=Pt(12)
#     ph_format.line_spacing=Pt(19)
    
    table=doc.add_table(rows = 1,cols = 3)
    # 获取第一行的单元格列表对象
    hdr_cells=table.rows[0].cells
    # 为每一个单元格赋值
    # 注：值都要为字符串类型
    hdr_cells[0].text='姓名__________'
    hdr_cells[1].text='开始时间__________'
    hdr_cells[2].text='结束时间__________'
    # 为表格添加一行
    for i in range(0,questlist.__len__(),3):
        print(i)
        new_cells = table.add_row().cells
        new_cells[0].text=questlist[i]
        if questlist.__len__()>i+1:
            new_cells[1].text=questlist[i+1]
        if questlist.__len__()>i+2:
            new_cells[2].text=questlist[i+2]
    
    fileFullName=getFileName(fileseq)
        
        
    doc.save(fileFullName)


#if __name__ == '__main__':
    #doc=Document()
    #doc=Document(u'D:\\Program Files\\liClipseWorkspace\\a.docx')
    #genDoc(doc)
    #print(getFileName(fileseq))


