#coding:utf-8
'''
Created on 2020年2月14日

@author: HEWEI
'''
from tkinter import *
import com.my.test.testDoc as gen
from docx import Document
from docx.shared import Inches,Pt
from docx.oxml.ns import qn
import random,os


fileseq=0

def chg_font(obj,fontname='微软雅黑',size=None):
    obj.font.name=fontname
    obj._element.rPr.rFonts.set(qn('w:eastAsia'),fontname)
    if size and isinstance(size, Pt):
        obj.font.size=size
        
def checkAndAddFormula(inta,symbol,intb,questline,questlist):
    if symbol=='+':
        questlist.append(questline)
    if symbol=='-':
        if inta>intb:
            questlist.append(questline)
    if symbol=='×':
        if inta!=0 and intb!=0:
            questlist.append(questline)
    if symbol=='÷':
        if inta!=0 and intb!=0 and inta%intb==0:
            questlist.append(questline)
    
        
def genContent(checkv,start,end,num):#count=30,plus=True,minus=True,multiply=False,divide=False):
    allsymbol=['+', '-', '×','÷']
    selectedsymbol=[]
    questlist=[]
    if checkv[0].get()==1:
        selectedsymbol.append(allsymbol[0])
    if checkv[1].get()==1:
        selectedsymbol.append(allsymbol[1])
    if checkv[2].get()==1:
        selectedsymbol.append(allsymbol[2])
    if checkv[3].get()==1:
        selectedsymbol.append(allsymbol[3])
    while True:
        inta=random.randint(start,end)
        stra=str(inta)
        symbol=random.choice(selectedsymbol)
        intb=random.randint(start,end)
        strb=str(intb)
        questline=stra+symbol+strb+'='
        checkAndAddFormula(inta,symbol,intb,questline,questlist)
        if questlist.__len__()==num:
            break
    for s in questlist:
        print(s)
    return questlist

def getFileName(fileseq):
    rootpath='D:\\数学\\'
    if not os.path.exists(rootpath):
        os.mkdir(rootpath)
        
    filename=rootpath+'a'+str(fileseq)+'.docx'
    if os.path.exists(filename):
        filename=getFileName(fileseq+1)
    return filename
        
def genDoc(checkv,start,end,num):
    doc=Document()
    distance=Inches(1)
    sec=doc.sections[0]
    sec.left_margin=distance
    sec.right_margin=distance
    sec.top_margin=distance
    sec.bottom_margin=distance
    sec.page_width=Inches(8.27)
    sec.page_height=Inches(11.75)
    
    chg_font(doc.styles['Normal'], fontname='宋体',size=Pt(14))
    
    questlist=genContent(checkv,start,end,num)
    
#     paragraph=doc.add_paragraph('姓名__________    开始时间__________    结束时间__________')
#     ph_format=paragraph.paragraph_format
#     ph_format.space_before=Pt(10)
#     ph_format.space_after=Pt(12)
#     ph_format.line_spacing=Pt(19)
    
    table=doc.add_table(rows = 1,cols = 3)
    # 获取第一行的单元格列表对象
    hdr_cells=table.rows[0].cells
    # 为每一个单元格赋值
    # 注：值都要为字符串类型
    hdr_cells[0].text='姓名__________'
    hdr_cells[1].text='开始时间__________'
    hdr_cells[2].text='结束时间__________'
    # 为表格添加一行
    for i in range(0,questlist.__len__(),3):
        new_cells = table.add_row().cells
        new_cells[0].text=questlist[i]
        if questlist.__len__()>i+1:
            new_cells[1].text=questlist[i+1]
        if questlist.__len__()>i+2:
            new_cells[2].text=questlist[i+2]
    
    fileFullName=getFileName(fileseq)
        
    doc.save(fileFullName)

def startGen(checkv,v1,v2,v3):
    for v in checkv:
        print(v.get())
#     v1=e1.getint()
#     v2=e2.getint()
#     v3=e3.getint()
    v1=v1.get()
    v2=v2.get()
    v3=v3.get()
    print(v1,v2,v3)
    genDoc(checkv,v1,v2,v3)
    var.set('已生成至"D:\数学"目录下')

root=Tk()
root.title('数学作业生成器')
# root.minsize(560, 545)  
# root.maxsize(560, 545) 

# class APP:
    
#     def __init__(self,master): #root传参赋值给master 啥意思?
#         frame=tk.Frame(master) #frame组件
#         #frame.pack(side=tk.LEFT,padx=10,pady=10)
#         frame.grid(row=0, column=0)
#         
#         self.hi_there=tk.Button(frame,text='打招呼',bg='black',fg='white',command=self.say_hi) #button按钮,command中调用定义的方法
#         #self.hi_there.pack()
#         self.hi_there.grid(row=0, column=1)
#         
#     def say_hi(self):
#         print('11111111111')
        
# frame=Frame(root)
# frame.pack(padx=10,pady=10)

# v=IntVar()
# c=Checkbutton(root,text='dddd',variable=v)
# c.pack()
# l=Label(root,textvariable=v)
# l.pack()

checkv=[]
SYMBOLS=['+', '-', '×','÷']
i=0
for symbol in SYMBOLS:
    checkv.append(IntVar())
    b=Checkbutton(root,text=symbol,variable=checkv[-1])
    #b.pack(anchor=W)
    b.grid(row=1, column=i)
    i=i+1
    
# for i in range(0,4):
#     checkv.append(IntVar())
#     b=Checkbutton(root,text=SYMBOLS[i],variable=checkv[-1])
#     #b.pack(anchor=W)
#     b.grid(row=1, column=i)
    
v1 = IntVar()
v2 = IntVar()
v3 = IntVar()

Label(root, text='开始:').grid(row=2, column=0)
e1 = Entry(root,width=10,textvariable=v1)
#e1.insert(10, 1)
e1.grid(row=2, column=1, padx=10, pady=5)

Label(root, text='结束:').grid(row=2, column=2)
e2 = Entry(root,width=10,textvariable=v2)
#e2.insert(10, 10)
e2.grid(row=2, column=3, padx=10, pady=5)

Label(root, text='数量:').grid(row=2, column=4)
e3 = Entry(root,width=10,textvariable=v3)
#e3.insert(10, 30)
e3.grid(row=2, column=5, padx=10, pady=5)

var=StringVar()
textLabel = Label(root,textvariable=var,justify=CENTER) 
textLabel.grid(row=3, column=0, padx=40, pady=5,columnspan=5)

Button(root,text='生成',command=lambda:startGen(checkv,v1,v2,v3)).grid(row=4,column=2,pady=5)

mainloop()